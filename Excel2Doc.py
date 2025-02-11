# This script reads the Excel file we use to track CRDC language corrections
# and converts it to a more readable report format pioneered by Mark C.
import docx.shared
import pandas as pd
import argparse
from crdclib import crdclib
import docx
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn



def readExcel(filename, sheetname):
    exeldf = pd.read_excel(filename, sheet_name=sheetname)
    return exeldf

def setCellBackground(cell, cellcolor):
    tablecell = cell._tc
    tablecellprops = tablecell.get_or_add_tcPr()
    cellshading = OxmlElement('w:shd')
    cellshading.set(qn('w:fill'), cellcolor)
    tablecellprops.append(cellshading)
    return cell


def setColumnWidth(column, width):
    for cell in column.cells:
        cell.width = width



def main(args):
    configs = crdclib.readYAML(args.configfile)
        
    #Create the document
    report = docx.Document()
    #Set font 
    styles = report.styles['Normal']
    font = styles.font
    font.name = 'Aptos Narrow'
    font.size = Pt(11)
    
    #Add the report title
    report.add_heading(configs['report_title'], 2)
    
    #Set up the table
    table = report.add_table(rows=1, cols=len(configs['table_headers']))
    table.style = 'Table Grid'
    
    #Create the header rows and content
    hdr_cells = table.rows[0].cells
  
    for i, header in enumerate(configs['table_headers']):
        hdr_cells[i].paragraphs[0].add_run(header).bold=True
        i = i+1
    j=0
    while j < len(configs['table_headers']):
       setCellBackground(hdr_cells[j], configs['headercolor'])
       j += 1
    
    # Put the header on every page
    table_header = OxmlElement('w:tblHeader')
    first_row_props = table.rows[0]._element.get_or_add_trPr()
    first_row_props.append(table_header)
    
    #Iterate through the workbooks
    #Linefill is used to color alternate rows light grey
    linefill = True
    for workbook in configs['workbooklist']:
        #Put the workbook in a df
        excel_df = readExcel(configs['excelfile'], workbook)
        #Remove empty rows
        excel_df.dropna(how='all', inplace=True)
        #Convert NaN
        excel_df = excel_df.fillna('')
        
        # Add a row announcing the new section, named afer the workbook
        # and colored blue
        newsec = table.add_row().cells
        newsec_cell = newsec[0]
        newsec_cell.paragraphs[0].add_run(f"{workbook}").bold=True
        setCellBackground(newsec[0], configs['sectioncolor'])
        setCellBackground(newsec[1], configs['sectioncolor'])
        
        #Iterate through the dataframe and populate the rows
        for index, row in excel_df.iterrows():
            cells = table.add_row().cells
            for entry in configs['content']:
                for cellindex, contentlist in entry.items():
                    for content in contentlist:
                        for exlabel, tablelabel in content.items():
                            textcontent = row[exlabel]
                            if textcontent == 0:
                                textcontent = ''
                            if tablelabel == 'None':
                                cells[cellindex].paragraphs[0].add_run(textcontent)
                                if linefill:
                                    setCellBackground(cells[cellindex], configs['linecolor'])
                            else:
                                workingcell = cells[cellindex]
                                workingcell.paragraphs[0].add_run(tablelabel).bold=True
                                workingcell.paragraphs[0].add_run(f"{row[exlabel]}\n")
                                if linefill:
                                    setCellBackground(workingcell, configs['linecolor'])
            if linefill:
                linefill = False
            else:
                linefill = True


    # With everything done, set the column widths
    table.autofit = False
    table.allow_autofit = False
    for entry in configs['colwidths']:
        for colindex, colwidth in entry.items():
            setColumnWidth(table.columns[colindex], docx.shared.Inches(colwidth))
    
    #Save the file
    report.save(configs['wordfile'])
        
        
    
    
    


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("-c", "--configfile", required=True,  help="Configuration file containing all the input info")

    args = parser.parse_args()

    main(args)